# tools.py
import plivo
import logging
import requests
from openai import OpenAI
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)
from google.cloud import storage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document
from openai import OpenAI
import os
from datetime import datetime
from io import StringIO
from dotenv import load_dotenv
import smtplib
from email.message import EmailMessage
from typing import List
import openai  # or use Gemini API depending on your stack
import os
from dotenv import load_dotenv
import re
import google.generativeai as genai 

import imaplib
import email
from email.header import decode_header
# calendar_utils.py
from datetime import datetime, timedelta
from datetime import datetime

import os
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

import os
from typing import Optional
from google.cloud import storage
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from openai import OpenAI
import fitz  # PyMuPDF
from docx import Document as DocxDocument

SCOPES = ['https://www.googleapis.com/auth/calendar']



load_dotenv()

sender_email = "rajawat.ajay@ehrlogic.com"
sender_password = "vkqv gpbc fxxq mgzx"



genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))  # Set your Google API key here


# Load environment variables from .env file
load_dotenv()

# Set the environment variable for Google Cloud credentials
gcs_credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if gcs_credentials_path:
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = gcs_credentials_path

def extract_text_from_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_and_query_rag_doc(file_path: Optional[str] = None, query: Optional[str] = None) -> str:
    try:
        if not file_path or not query:
            return "❌ Missing file path or query for RAG processing."

        # Step 1: Read file content
        raw_text = ""

        if file_path.startswith("gs://"):
            # Extract bucket and blob name
            bucket_name, blob_name = file_path.replace("gs://", "").split("/", 1)

            # Access GCS
            client = storage.Client()
            bucket = client.bucket(bucket_name)
            blob = bucket.blob(blob_name)

            # Download file temporarily
            temp_path = f"/tmp/{os.path.basename(blob_name)}"
            blob.download_to_filename(temp_path)
            file_path = temp_path

        # Extract text based on file type
        if file_path.endswith(".pdf"):
            raw_text = extract_text_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            raw_text = extract_text_from_docx(file_path)
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()

        if not raw_text.strip():
            return "⚠️ File is empty or could not be read."

        # Step 2: Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
        texts = text_splitter.split_text(raw_text)
        docs = [Document(page_content=t) for t in texts]

        # Step 3: Create vector DB
        embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY"))
        vectordb = FAISS.from_documents(docs, embeddings)

        # Step 4: Retrieve relevant context
        context_docs = vectordb.as_retriever().get_relevant_documents(query)
        context_text = "\n\n".join([doc.page_content for doc in context_docs[:3]])

        # Step 5: Query LLM
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a RAG Agent tool. Follow the user's query using the given document content."},
                {"role": "user", "content": f"Context:\n{context_text}\n\nQuestion: {query}"}
            ],
            temperature=0.3
        )

        answer = completion.choices[0].message.content.strip()
        return answer

    except Exception as e:
        return f"❌ RAG failed: {str(e)}"


def request_call(phone_number: str) -> str:
    try:
        phone_number = phone_number.strip().replace('"', '')
        AUTH_ID = "MAZDC0M2EYNJZKZTC4NT"
        AUTH_TOKEN = "ZmIyNTU2YTlkZWIzNmEzNmMwYzdmZTZlMDdlMGRh"

        client = plivo.RestClient(AUTH_ID, AUTH_TOKEN)
        call_made = client.calls.create(
            from_='+912269976361',
            to_=phone_number,
            answer_url='https://viable-firstly-cicada.ngrok-free.app/webhook',
            answer_method='GET'
        )
        return "Call initiated successfully"
    except Exception as e:
        logger.error(f"❌ Error: {str(e)}")
        return f"Call initiation failed: {str(e)}"


def join_meeting(meeting_url: str, bot_name: str) -> str:
    try:
        response = requests.post(
            url="https://us-west-2.recall.ai/api/v1/bot/",
            headers={
                "accept": "application/json",
                "content-type": "application/json",
                "Authorization": "7afe842f24e6129cd72899cb83ed4e6757a84401"
            },
            json={
                "recording_config": {
                    "transcript": {
                        "provider": {
                            "meeting_captions": {}
                        }
                    },
                    "realtime_endpoints": [
                        {
                            "type": "webhook",
                            "url": "https://viable-firstly-cicada.ngrok-free.app/recall/transcript",
                            "events": ["transcript.data", "transcript.partial_data", "participant_events.chat_message"]
                        }
                    ]
                },
                "meeting_url": meeting_url,
                "bot_name": bot_name,
                "chat": {
                    "on_bot_join": {
                        "send_to": "everyone",
                        "message": f"Hey all! I'm {bot_name}, your AI assistant. I'll be taking notes!"
                    }
                }
            }
        )
        result = response.json()
        bot_id = result.get("id", "No ID returned")
        return f"Bot '{bot_name}' joined meeting. Bot ID: {bot_id}"
    except Exception as e:
        return f"❌ Failed to join meeting: {str(e)}"



def fetch_latest_email(username: str, password: str) -> str:
    try:
        imap_server = "imap.gmail.com"

        # Connect and login
        mail = imaplib.IMAP4_SSL(imap_server)
        mail.login(username, password)
        mail.select("inbox")

        status, messages = mail.search(None, "ALL")
        email_ids = messages[0].split()

        if not email_ids:
            return "❌ No emails found in the inbox."

        # Fetch the latest email
        email_id = email_ids[-1]
        status, msg_data = mail.fetch(email_id, "(RFC822)")

        for response_part in msg_data:
            if isinstance(response_part, tuple):
                msg = email.message_from_bytes(response_part[1])
                subject, encoding = decode_header(msg["Subject"])[0]
                subject = subject.decode(encoding if encoding else "utf-8") if isinstance(subject, bytes) else subject
                from_ = msg.get("From")
                date = msg.get("Date")

                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        content_type = part.get_content_type()
                        if content_type == "text/plain":
                            body = part.get_payload(decode=True).decode()
                            break
                else:
                    body = msg.get_payload(decode=True).decode()

                return f"From: {from_}\nSubject: {subject}\nDate: {date}\nBody:\n{body.strip()}"


        return "❌ Unable to read the email content."

    except Exception as e:
        return f"❌ Error fetching email: {str(e)}"



def get_calendar_service():
    creds = None
    token_path = 'token.json'
    creds_path = 'credentials.json'  # Path to your OAuth file

    if os.path.exists(token_path):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    else:
        flow = InstalledAppFlow.from_client_secrets_file(creds_path, SCOPES)
        creds = flow.run_local_server(port=0)
        with open(token_path, 'w') as token:
            token.write(creds.to_json())

    service = build('calendar', 'v3', credentials=creds)
    return service

def schedule_meeting(start_time: datetime, emails: list, duration_minutes=30, summary="Team Meeting"):
    service = get_calendar_service()

    end_time = start_time + timedelta(minutes=duration_minutes)
    event = {
        'summary': summary,
        'start': {'dateTime': start_time.isoformat(), 'timeZone': 'Asia/Kolkata'},
        'end': {'dateTime': end_time.isoformat(), 'timeZone': 'Asia/Kolkata'},
        'conferenceData': {
            'createRequest': {
                'conferenceSolutionKey': {'type': 'hangoutsMeet'},
                'requestId': 'virtus-meeting-' + start_time.isoformat()
            }
        },
        'attendees': [{'email': email} for email in emails]
    }

    event = service.events().insert(
        calendarId='primary',
        body=event,
        conferenceDataVersion=1,
        sendUpdates="all"
    ).execute()

    return event.get("htmlLink")


def generate_email_from_prompt(message: str, sender_name: str = "Ajay Singh") -> tuple:
    """Use Gemini LLM to generate a professional subject and formatted email body from a casual message."""
    prompt = f"""
You are a professional communication assistant.
Take the following casual or natural message and do three things:
1. Generate a professional email subject.
2. Convert the message into a clean, polite, formal email body.
3. Add a professional closing line with 'Regards, {sender_name}' at the end.

Message:
\"\"\"{message}\"\"\"
Please respond in the following format:
Subject: <subject line>
<email body>
"""

    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(prompt)

    output = response.text.strip()

    # Extract subject and body
    subject_match = re.search(r"Subject:\s*(.*)", output)
    subject = subject_match.group(1).strip() if subject_match else "No Subject"
    body = re.sub(r"Subject:.*", "", output, flags=re.DOTALL).strip()

    return subject


def smart_email_sender(recipients: List[str], message: str, sender_name: str = "Ajay"):
    """
    Sends a professional email based on a natural language input.
    """
    subject = generate_email_from_prompt(message, sender_name)

    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = sender_email
    msg["To"] = ", ".join(recipients)
    msg.set_content(message)

    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(sender_email, sender_password)
            smtp.send_message(msg)
        return f"Email sent successfully to {', '.join(recipients)}."
    except Exception as e:
        return {"status": "error", "message": str(e)}
    

